"""
Export lesson plans to teacher-friendly formats (TXT, DOCX).
Teachers can actually use these files!
"""

from datetime import datetime
from typing import Dict, Any


def export_to_txt(lesson_plan: Dict[str, Any], filename: str = None) -> str:
    """
    Export lesson plan to a readable TXT file.
    
    Args:
        lesson_plan: The lesson plan dictionary
        filename: Optional filename, auto-generated if not provided
        
    Returns:
        Path to the created file
    """
    metadata = lesson_plan.get('metadata', {})
    
    # Generate filename if not provided
    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        grade = metadata.get('grade', 'X')
        subject = metadata.get('subject', 'lesson').replace(' ', '_')
        filename = f"Урок_{grade}_клас_{subject}_{timestamp}.txt"
    
    # Build the content
    lines = []
    lines.append("=" * 80)
    lines.append("ПЛАН УРОКУ")
    lines.append("Згенеровано системою Pedagogue AI")
    lines.append("=" * 80)
    lines.append("")
    
    # Metadata
    lines.append("ІНФОРМАЦІЯ ПРО УРОК")
    lines.append("-" * 80)
    lines.append(f"Клас:           {metadata.get('grade')}")
    lines.append(f"Предмет:        {metadata.get('subject')}")
    lines.append(f"Тема:           {metadata.get('topic')}")
    lines.append(f"Тривалість:     {metadata.get('duration')} хвилин")
    lines.append(f"Дата створення: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    lines.append("")
    
    # Quality score
    qa_review = lesson_plan.get('qa_review', {})
    if qa_review.get('success'):
        scores = qa_review.get('scores', {})
        lines.append("ОЦІНКА ЯКОСТІ")
        lines.append("-" * 80)
        lines.append(f"Загальна оцінка: {scores.get('average', 0):.1f}/10")
        lines.append(f"Статус:          {qa_review.get('overall_status', 'unknown')}")
        lines.append("")
    
    # НУШ Standards (if available)
    standards = lesson_plan.get('standards', {})
    if standards.get('success'):
        lines.append("СТАНДАРТИ НУШ")
        lines.append("-" * 80)
        lines.append(standards.get('standards', 'Не знайдено'))
        lines.append("")
    
    # Learning Strategies
    strategies = lesson_plan.get('learning_strategies', {})
    if strategies.get('success'):
        lines.append("СТРАТЕГІЇ НАВЧАННЯ")
        lines.append("-" * 80)
        lines.append(strategies.get('strategies', ''))
        lines.append("")
    
    # Learning Objectives
    objectives = lesson_plan.get('objectives', {})
    if objectives.get('success'):
        lines.append("ЦІЛІ НАВЧАННЯ (SMART)")
        lines.append("-" * 80)
        lines.append(objectives.get('objectives', ''))
        lines.append("")
    
    # Warm-up Activity
    warmup = lesson_plan.get('warmup', {})
    if warmup.get('success'):
        lines.append("РОЗМИНКА (5-10 хвилин)")
        lines.append("-" * 80)
        lines.append(warmup.get('warmup', ''))
        lines.append("")
    
    # Direct Instruction
    instruction = lesson_plan.get('instruction', {})
    if instruction.get('success'):
        lines.append("ПРЯМЕ НАВЧАННЯ")
        lines.append("-" * 80)
        lines.append(instruction.get('instruction', ''))
        lines.append("")
    
    # Guided Practice
    guided = lesson_plan.get('guided_practice', {})
    if guided.get('success'):
        lines.append("КЕРОВАНА ПРАКТИКА")
        lines.append("-" * 80)
        lines.append(guided.get('activity', ''))
        lines.append("")
    
    # Differentiation
    diff = lesson_plan.get('differentiation', {})
    if diff.get('success'):
        lines.append("ДИФЕРЕНЦІАЦІЯ (3 РІВНІ)")
        lines.append("-" * 80)
        
        tiers = diff.get('tiers', {})
        if 'базовий' in tiers:
            lines.append("\n[БАЗОВИЙ РІВЕНЬ]")
            lines.append(tiers['базовий'])
        if 'середній' in tiers:
            lines.append("\n[СЕРЕДНІЙ РІВЕНЬ]")
            lines.append(tiers['середній'])
        if 'високий' in tiers:
            lines.append("\n[ВИСОКИЙ РІВЕНЬ]")
            lines.append(tiers['високий'])
        lines.append("")
    
    # Independent Practice
    independent = lesson_plan.get('independent_practice', {})
    if independent.get('success'):
        lines.append("САМОСТІЙНА ПРАКТИКА")
        lines.append("-" * 80)
        lines.append(independent.get('activity', ''))
        lines.append("")
    
    # Formative Assessment
    formative = lesson_plan.get('formative_assessment', {})
    if formative.get('success'):
        lines.append("ФОРМУВАЛЬНЕ ОЦІНЮВАННЯ")
        lines.append("-" * 80)
        lines.append(formative.get('items', ''))
        lines.append("")
    
    # Summative Assessment
    summative = lesson_plan.get('summative_assessment', {})
    if summative.get('success'):
        lines.append("ПІДСУМКОВЕ ОЦІНЮВАННЯ")
        lines.append("-" * 80)
        lines.append(summative.get('assessment', ''))
        lines.append("")
    
    # QA Review and Suggestions
    if qa_review.get('success'):
        lines.append("РЕКОМЕНДАЦІЇ ЕКСПЕРТА")
        lines.append("-" * 80)
        review_text = qa_review.get('review', '')
        # Ensure review is string
        if isinstance(review_text, str):
            lines.append(review_text)
        else:
            lines.append(str(review_text))
        
        suggestions = qa_review.get('suggestions', [])
        if suggestions:
            lines.append("\nПропозиції для покращення:")
            for i, suggestion in enumerate(suggestions, 1):
                # Ensure suggestion is string
                if isinstance(suggestion, str):
                    lines.append(f"{i}. {suggestion}")
                else:
                    lines.append(f"{i}. {str(suggestion)}")
        lines.append("")
    
    # Footer
    lines.append("=" * 80)
    lines.append("Кінець плану уроку")
    lines.append("Згенеровано Pedagogue AI - https://github.com/your-repo")
    lines.append("=" * 80)
    
    # Write to file - ensure all lines are strings
    try:
        content = "\n".join(str(line) for line in lines)
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(content)
    except Exception as e:
        print(f"[ERROR] Failed to write file: {e}")
        # Try again with more aggressive string conversion
        content = "\n".join(str(line) if not isinstance(line, str) else line for line in lines)
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(content)
    
    return filename


def export_to_docx(lesson_plan: Dict[str, Any], filename: str = None) -> str:
    """
    Export lesson plan to a formatted DOCX file.
    Requires python-docx package.
    
    Args:
        lesson_plan: The lesson plan dictionary
        filename: Optional filename, auto-generated if not provided
        
    Returns:
        Path to the created file
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        # Fallback to TXT if docx not available
        print("[!] python-docx not installed. Saving as TXT instead.")
        print("[!] To enable DOCX: pip install python-docx")
        return export_to_txt(lesson_plan, filename.replace('.docx', '.txt') if filename else None)
    
    metadata = lesson_plan.get('metadata', {})
    
    # Generate filename if not provided
    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        grade = metadata.get('grade', 'X')
        subject = metadata.get('subject', 'lesson').replace(' ', '_')
        filename = f"Урок_{grade}_клас_{subject}_{timestamp}.docx"
    
    # Create document
    doc = Document()
    
    # Title
    title = doc.add_heading('ПЛАН УРОКУ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Згенеровано системою Pedagogue AI')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(10)
    subtitle.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # Metadata table
    doc.add_heading('Інформація про урок', 1)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    table.cell(0, 0).text = 'Клас:'
    table.cell(0, 1).text = str(metadata.get('grade', ''))
    table.cell(1, 0).text = 'Предмет:'
    table.cell(1, 1).text = metadata.get('subject', '')
    table.cell(2, 0).text = 'Тема:'
    table.cell(2, 1).text = metadata.get('topic', '')
    table.cell(3, 0).text = 'Тривалість:'
    table.cell(3, 1).text = f"{metadata.get('duration', '')} хвилин"
    table.cell(4, 0).text = 'Дата створення:'
    table.cell(4, 1).text = datetime.now().strftime('%d.%m.%Y %H:%M')
    
    # Quality score
    qa_review = lesson_plan.get('qa_review', {})
    if qa_review.get('success'):
        doc.add_heading('Оцінка якості', 1)
        scores = qa_review.get('scores', {})
        p = doc.add_paragraph()
        p.add_run(f"Загальна оцінка: ").bold = True
        p.add_run(f"{scores.get('average', 0):.1f}/10\n")
        p.add_run("Статус: ").bold = True
        p.add_run(qa_review.get('overall_status', 'unknown'))
    
    # Add sections
    sections = [
        ('Стандарти НУШ', 'standards', 'standards'),
        ('Стратегії навчання', 'learning_strategies', 'strategies'),
        ('Цілі навчання (SMART)', 'objectives', 'objectives'),
        ('Розминка (5-10 хвилин)', 'warmup', 'warmup'),
        ('Пряме навчання', 'instruction', 'instruction'),
        ('Керована практика', 'guided_practice', 'activity'),
        ('Самостійна практика', 'independent_practice', 'activity'),
        ('Формувальне оцінювання', 'formative_assessment', 'items'),
        ('Підсумкове оцінювання', 'summative_assessment', 'assessment'),
    ]
    
    for title, key, content_key in sections:
        section = lesson_plan.get(key, {})
        if section.get('success'):
            doc.add_heading(title, 1)
            doc.add_paragraph(section.get(content_key, ''))
    
    # Differentiation (special handling)
    diff = lesson_plan.get('differentiation', {})
    if diff.get('success'):
        doc.add_heading('Диференціація (3 рівні)', 1)
        tiers = diff.get('tiers', {})
        
        if 'базовий' in tiers:
            doc.add_heading('Базовий рівень', 2)
            doc.add_paragraph(tiers['базовий'])
        if 'середній' in tiers:
            doc.add_heading('Середній рівень', 2)
            doc.add_paragraph(tiers['середній'])
        if 'високий' in tiers:
            doc.add_heading('Високий рівень', 2)
            doc.add_paragraph(tiers['високий'])
    
    # QA Review
    if qa_review.get('success'):
        doc.add_heading('Рекомендації експерта', 1)
        doc.add_paragraph(qa_review.get('review', ''))
        
        suggestions = qa_review.get('suggestions', [])
        if suggestions:
            doc.add_heading('Пропозиції для покращення', 2)
            for suggestion in suggestions:
                doc.add_paragraph(suggestion, style='List Bullet')
    
    # Footer
    doc.add_page_break()
    footer = doc.add_paragraph('Згенеровано Pedagogue AI')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # Save
    doc.save(filename)
    
    return filename
